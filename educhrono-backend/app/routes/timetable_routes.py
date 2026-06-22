from fastapi import APIRouter
from fastapi.responses import FileResponse
from app import db
from app.services.timetable_service import generate_timetable
from fpdf import FPDF
from datetime import datetime
from docx import Document
import os
import tempfile

router = APIRouter()

# ====================================================================================
# ✅ 1. GENERATE TIMETABLE ROUTE (FIXED - NO 500 ERROR)
# ====================================================================================
@router.post("/generate")
async def generate_timetable_api():
    try:
        result = generate_timetable()

        return {
            "success": True,
            "message": "✅ Timetable generated successfully",
            "total_entries": result.get("count", 0)
        }

    except Exception as e:
        print("🔥 ERROR:", e)   # 👈 IMPORTANT DEBUG
        return {
            "success": False,
            "error": str(e)
        }


# ====================================================================================
# ✅ 2. FETCH COMPLETE TIMETABLE
# ====================================================================================
@router.get("/list")
async def fetch_timetable():
    try:
        data = list(db["timetable"].find({}, {"_id": 0}))
        return {"success": True, "total": len(data), "timetable": data}

    except Exception as e:
        print("🔥 ERROR:", e)
        return {"success": False, "error": str(e)}


# ====================================================================================
# ✅ 3. CLEAR TIMETABLE
# ====================================================================================
@router.delete("/clear")
async def clear_timetable():
    try:
        db["timetable"].delete_many({})
        return {"success": True, "message": "Timetable cleared."}

    except Exception as e:
        print("🔥 ERROR:", e)
        return {"success": False, "error": str(e)}


# ====================================================================================
# ✅ 4. FACULTY TIMETABLE
# ====================================================================================
@router.get("/faculty/{faculty_name}")
def get_faculty_timetable(faculty_name: str):

    timetable = list(db["timetable"].find({}, {"_id": 0}))
    faculty_entries = []

    for entry in timetable:

        if entry.get("Faculty") == faculty_name:
            faculty_entries.append(entry)
            continue

        if entry.get("Type") == "P" and "Groups" in entry:
            for g in entry.get("Groups", []):
                if g.get("faculty") == faculty_name:
                    faculty_entries.append({
                        "Semester": entry["Semester"],
                        "Section": entry["Section"],
                        "Day": entry["Day"],
                        "Type": "P",
                        "Slot": entry["Slot"],
                        "Slot2": entry.get("Slot2"),
                        "Subject": g.get("subject"),
                        "Faculty": g.get("faculty"),
                        "Room": g.get("room")
                    })
                    break

    return {
        "success": True,
        "faculty": faculty_name,
        "total": len(faculty_entries),
        "timetable": faculty_entries
    }


# ====================================================================================
# ✅ 5. STUDENT TIMETABLE
# ====================================================================================
@router.get("/student/{semester}/{section}")
def get_student_timetable(semester: int, section: str):

    data = list(db["timetable"].find(
        {"Semester": semester, "Section": section}, {"_id": 0}
    ))

    return {
        "success": True,
        "semester": semester,
        "section": section,
        "total": len(data),
        "timetable": data
    }


# ====================================================================================
# ✅ 6. DOWNLOAD PDF
# ====================================================================================
@router.get("/download/{role}/{identifier}")
def download_timetable(role: str, identifier: str):

    filename = ""
    data = []
    timetable = list(db["timetable"].find({}, {"_id": 0}))

    if role == "admin":
        data = timetable
        filename = "Full_Timetable.pdf"

    elif role == "faculty":
        for entry in timetable:

            if entry.get("Faculty") == identifier:
                data.append(entry)
                continue

            if entry.get("Type") == "P" and "Groups" in entry:
                for g in entry.get("Groups", []):
                    if g.get("faculty") == identifier:
                        data.append({
                            "Day": entry["Day"],
                            "Slot": entry["Slot"],
                            "Subject": g.get("subject"),
                            "Faculty": g.get("faculty"),
                            "Room": g.get("room")
                        })
                        break

        filename = f"Faculty_{identifier.replace(' ', '_')}.pdf"

    elif role == "student":
        sem, section = identifier.split("-")
        data = list(db["timetable"].find(
            {"Semester": int(sem), "Section": section}, {"_id": 0}
        ))
        filename = f"Student_Sem{sem}_Sec{section}.pdf"

    else:
        return {"success": False, "error": "Invalid role"}

    if not data:
        return {"success": False, "error": "No timetable found"}

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
    slots = [
        "09:00-10:00",
        "10:00-11:00",
        "11:00-12:00",
        "12:00-13:00",
        "14:00-15:00",
        "15:00-16:00",
        "16:00-17:00"
    ]

    grid = {day: {slot: "" for slot in slots} for day in days}

    for entry in data:
        grid[entry["Day"]][entry["Slot"]] = f"{entry.get('Subject','')}\n{entry.get('Faculty','')}\n({entry.get('Room','')})"

    pdf = FPDF("L", "mm", "A4")
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 15)
    pdf.cell(0, 10, "EduChrono Timetable", ln=True, align="C")

    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Generated on: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}", ln=True, align="C")

    col_width = 260 / len(slots)
    row_height = 10

    pdf.cell(25, row_height, "Day", border=1)

    for slot in slots:
        pdf.cell(col_width, row_height, slot, border=1)

    pdf.ln(row_height)

    for day in days:
        pdf.cell(25, row_height * 3, day, border=1)

        for slot in slots:
            text = grid[day][slot] or "-"
            x = pdf.get_x()
            y = pdf.get_y()
            pdf.multi_cell(col_width, 5, text, border=1)
            pdf.set_xy(x + col_width, y)

        pdf.ln(row_height * 3)

    path = os.path.join(tempfile.gettempdir(), filename)
    pdf.output(path)

    return FileResponse(path, filename=filename)

@router.get("/download-word/{role}/{identifier}")
def download_timetable_word(role: str, identifier: str):

    role = role.lower().strip()

    filename = ""
    data = []
    timetable = list(db["timetable"].find({}, {"_id": 0}))

    # 🔹 SAME FILTER (PDF jaisa)
    if role == "admin":
        data = timetable
        filename = "Full_Timetable.docx"

    elif role == "faculty":
        for entry in timetable:

            if entry.get("Faculty") == identifier:
                data.append(entry)
                continue

            if entry.get("Type") == "P" and "Groups" in entry:
                for g in entry.get("Groups", []):
                    if g.get("faculty") == identifier:
                        data.append({
                            "Day": entry["Day"],
                            "Slot": entry["Slot"],
                            "Subject": g.get("subject"),
                            "Faculty": g.get("faculty"),
                            "Room": g.get("room")
                        })
                        break

        filename = f"Faculty_{identifier.replace(' ', '_')}.docx"

    elif role == "student":
        sem, section = identifier.split("-")
        data = list(db["timetable"].find(
            {"Semester": int(sem), "Section": section}, {"_id": 0}
        ))
        filename = f"Student_Sem{sem}_Sec{section}.docx"

    else:
        return {"success": False, "error": "Invalid role"}

    if not data:
        return {"success": False, "error": "No timetable found"}

    # 🔥 SAME GRID AS PDF (NO CHANGE)
    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
    slots = [
        "09:00-10:00",
        "10:00-11:00",
        "11:00-12:00",
        "12:00-13:00",
        "14:00-15:00",
        "15:00-16:00",
        "16:00-17:00"
    ]

    grid = {day: {slot: "" for slot in slots} for day in days}

    for entry in data:
        grid[entry["Day"]][entry["Slot"]] = (
            f"{entry.get('Subject','')}\n"
            f"{entry.get('Faculty','')}\n"
            f"({entry.get('Room','')})"
        )

    # 🔥 WORD TABLE (ONLY RENDERING CHANGE)
    from docx import Document
    doc = Document()
    doc.add_heading("EduChrono Timetable", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}")

    table = doc.add_table(rows=len(days)+1, cols=len(slots)+1)
    table.style = "Table Grid"

    # HEADER
    table.cell(0, 0).text = "Day"
    for j, slot in enumerate(slots):
        table.cell(0, j+1).text = slot

    # DATA (same text)
    for i, day in enumerate(days):
        table.cell(i+1, 0).text = day
        for j, slot in enumerate(slots):
            text = grid[day][slot] or "-"
            table.cell(i+1, j+1).text = text

    path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(path)

    return FileResponse(
        path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



